"""
Phase 2 polish:
- Algorithm 4-8 pseudocode: replace "生成资金解冻消息" and "生成股票解冻消息"
  with HTTP call expressions, consistent with sync HTTP design.
- Add a clarifying note in 4.2.4.3 (账户结算) introduction.
"""
from docx import Document

SRC = "/Users/bob/WeChatProjects/skihive_v1/.reports/系统设计报告.modified.docx"
DST = "/Users/bob/WeChatProjects/skihive_v1/.reports/系统设计报告.modified.docx"

doc = Document(SRC)

# Find and fix pseudocode 4-8
fixed = 0
for i, p in enumerate(doc.paragraphs):
    if p.style.name != "伪代码":
        continue
    if "生成资金解冻消息" in p.text:
        p.text = p.text.replace(
            "生成资金解冻消息 (order.order_id, order.fund_account_id, release_amount);",
            "// 同步调用 ACCOUNT 资金解冻接口（POST /fund-accounts/{id}/release），"
            "按响应结果推进；message_id 用于幂等去重。"
        )
        fixed += 1
    if "生成股票解冻消息" in p.text:
        p.text = p.text.replace(
            "生成股票解冻消息 (order.order_id, order.security_account_id, order.stock_code, release_quantity);",
            "// 同步调用 ACCOUNT 证券解冻接口（POST /security-accounts/{id}/positions/release），"
            "按响应结果推进；message_id 用于幂等去重。"
        )
        fixed += 1
print(f"[OK] Updated {fixed} pseudocode lines in 算法 4-8")

# Update 4.2.4.3 账户结算 module overview (paragraph 498) to make HTTP sync explicit
def find_para_index(text_substr, style_name=None):
    for i, p in enumerate(doc.paragraphs):
        if text_substr in p.text:
            if style_name is None or p.style.name == style_name:
                return i
    return -1

idx = find_para_index("成交后调用 ACCOUNT 系统接口完成资金和证券变更")
if idx >= 0:
    doc.paragraphs[idx].text = (
        "成交后 TRADE 通过 HTTP RESTful API（同步）调用 ACCOUNT 的资金结算、证券结算与冻结资源释放接口，"
        "等待 HTTP 响应并按返回状态推进后续流程；同步调用配合 message_id 幂等与重试机制保证一致性，"
        "本期不引入 MQ 异步解耦。"
    )
    print(f"[OK] Updated 4.2.4.3 账户结算 overview (paragraph {idx})")

doc.save(DST)
print(f"[SAVED] {DST}")
