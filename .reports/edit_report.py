"""
Optimize 系统设计报告.docx:
1. Remove stray 's' character
2. Update ACCOUNT communication: HTTP synchronous (MQ as future extension)
3. Update order book data structure description
4. Clean up obvious formatting issues (excess empty paragraphs on cover)
"""
from docx import Document
from docx.oxml.ns import qn as _qn  # not used but reserved

SRC = "/Users/bob/WeChatProjects/skihive_v1/.reports/系统设计报告.original.docx"
DST = "/Users/bob/WeChatProjects/skihive_v1/.reports/系统设计报告.modified.docx"

doc = Document(SRC)

# ---------- 1. Remove the stray 's' (paragraph 136) ----------
p = doc.paragraphs[136]
assert p.text.strip() == "s", f"Expected 's' at 136, got: {repr(p.text)}"
p_elem = p._element
p_elem.getparent().remove(p_elem)
print(f"[OK] Removed stray 's' at paragraph 136")

# ---------- 2. Update Table 16 (3.4.3 消息队列与通信) ----------
new_rows = [
    ("类别", "选型", "说明"),
    ("TRADE ↔ ACCOUNT 通信", "HTTP RESTful API（同步）",
     "本系统实现采用 HTTP RESTful API 作为 TRADE 与 ACCOUNT 的统一通信方式。冻结复核、结算、资源释放等请求均通过 POST 同步调用并等待响应结果，调用方根据返回状态决定后续流程。该方式与本系统接口文档以 HTTP 接口为主保持一致，联调成本低、可观测性强。"),
    ("服务间异步扩展（后续）", "RabbitMQ 3.10+（推荐），Kafka 3.0+（备选）",
     "本次实现不依赖 MQ 完成 TRADE ↔ ACCOUNT 通信。MQ 仅作为后续扩展方案保留：例如高吞吐场景下对结算事件、行情事件、客户端反馈事件的异步解耦。引入 MQ 时需同时设计消息确认、幂等消费、outbox 补偿等机制，本期不实现。"),
    ("反向代理 / 入口网关", "Nginx 1.22+",
     "作为入口网关，提供 TLS 终止与反向代理。TRADE 单实例部署，Nginx 仅做反向代理转发；INFO / ADMIN 可水平扩展，Nginx 对其做负载均衡。"),
]

t = doc.tables[16]
# Modify existing rows in place
for i, row in enumerate(t.rows):
    if i >= len(new_rows):
        break
    new_data = new_rows[i]
    for j, cell in enumerate(row.cells):
        # Clear cell paragraphs (keep first paragraph, remove the rest)
        cell_paras = list(cell.paragraphs)
        for p_in_cell in cell_paras[1:]:
            p_in_cell._element.getparent().remove(p_in_cell._element)
        if j < len(new_data):
            cell.text = new_data[j]
            if cell.paragraphs:
                # Keep Normal style for the cell text
                cell.paragraphs[0].style = doc.styles['Normal']

# Add rows if table has fewer than new_rows
while len(t.rows) < len(new_rows):
    new_row = t.add_row()
    idx = len(t.rows) - 1
    new_data = new_rows[idx]
    for j, cell in enumerate(new_row.cells):
        if j < len(new_data):
            cell.text = new_data[j]
            if cell.paragraphs:
                cell.paragraphs[0].style = doc.styles['Normal']

print(f"[OK] Updated Table 16 (3.4.3 消息队列与通信) -> {len(t.rows)} rows")

# ---------- 3. Update 3.4.6 #5 (异步 HTTP 客户端) ----------
def find_para_index(text_substr, style_name=None):
    for i, p in enumerate(doc.paragraphs):
        if text_substr in p.text:
            if style_name is None or p.style.name == style_name:
                return i
    return -1

idx_async_http = find_para_index("异步 HTTP 客户端")
assert idx_async_http >= 0, "Could not find 异步 HTTP 客户端"
new_text_5 = (
    "5. 同步 HTTP 客户端：使用 httpx Client 同步调用 ACCOUNT 模块的冻结复核、结算、资源释放接口。"
    "TRADE 在请求结束后同步等待响应结果，根据返回的成功/失败状态决定后续流程，保证关键资金/证券操作的事务一致性。"
    "同步调用配合明确的超时、重试和幂等机制，避免引入 MQ 带来的消息确认、重复消费与 outbox 补偿复杂度。"
)
doc.paragraphs[idx_async_http].text = new_text_5
print(f"[OK] Updated 3.4.6 #5 (paragraph {idx_async_http}): 异步→同步 HTTP")

# ---------- 4. Update 3.5.3 部署图说明 points 5, 6 ----------
idx_p5 = find_para_index("5. 消息队列异步解耦")
idx_p6 = find_para_index("6. 交易与结算系统与资金账户系统")
assert idx_p5 >= 0 and idx_p6 >= 0, "Could not find MQ deployment bullets"

doc.paragraphs[idx_p5].text = (
    "5. TRADE ↔ ACCOUNT 通信采用 HTTP RESTful API（同步）：TRADE 单实例应用服务器通过内部网络直接以同步 HTTP 调用 ACCOUNT 接口，"
    "包括冻结复核、结算、资源释放等关键路径；调用方按 HTTP 响应结果推进流程。MQ 在本期不作为 TRADE ↔ ACCOUNT 的必要通信方式，仅在后续扩展中按需引入。"
)
doc.paragraphs[idx_p6].text = (
    "6. Nginx 对 TRADE 仅做反向代理转发（单实例不参与负载均衡）；INFO / ADMIN 模块独立水平扩展，Nginx 对其做负载均衡。"
    "TRADE 内部各协程之间使用进程内直接调用（in-process），不依赖外部消息队列中转。"
)
print(f"[OK] Updated 3.5.3 部署图说明 points 5,6 (paragraphs {idx_p5}, {idx_p6})")

# ---------- 5. Update 7.2 启动流程 point 3 ----------
idx_mq_start = find_para_index("启动 RabbitMQ/Kafka 消息队列服务")
assert idx_mq_start >= 0, "Could not find RabbitMQ start"
doc.paragraphs[idx_mq_start].text = (
    "3. （可选 / 扩展）若后续启用 MQ 异步扩展，按需启动 RabbitMQ/Kafka 消息队列服务并创建相应队列；"
    "本期 TRADE ↔ ACCOUNT 通信以 HTTP 同步为主，无需 MQ。"
)
print(f"[OK] Updated 7.2 system startup point 3 (paragraph {idx_mq_start})")

# ---------- 6. Update 7.4 结算反馈协程 ----------
idx_settle_coro = find_para_index("结算反馈协程")
assert idx_settle_coro >= 0, "Could not find 结算反馈协程"
doc.paragraphs[idx_settle_coro].text = (
    "3. 结算反馈协程：撮合产出成交记录后，按成交逐笔调用 SettlementService，通过 httpx Client 同步调用 ACCOUNT 模块的资金结算、"
    "证券结算与冻结资源释放接口，等待响应并按返回状态推进后续流程（反馈推送、行情更新）。同步调用配合消息幂等和重试机制保证一致性。"
)
print(f"[OK] Updated 7.4 结算反馈协程 (paragraph {idx_settle_coro})")

# ---------- 7. Update 8.2.1 消息队列故障恢复 ----------
idx_mq_recovery = find_para_index("消息队列故障恢复")
assert idx_mq_recovery >= 0, "Could not find 消息队列故障恢复"

doc.paragraphs[idx_mq_recovery].text = "消息队列扩展方案（本期不实现）"
doc.paragraphs[idx_mq_recovery + 1].text = (
    "a. 后续若启用 MQ 异步扩展，建议采用 RabbitMQ/Kafka 集群部署：单节点故障不影响整体服务，消息自动路由至其他节点。"
)
doc.paragraphs[idx_mq_recovery + 2].text = (
    "b. 持久化消息（结算、反馈、释放事件）在 Broker 重启后从磁盘恢复，不会丢失；同时需设计消息幂等消费与 outbox 补偿机制，避免重复处理。"
)
print(f"[OK] Updated 8.2.1 消息队列故障恢复 (paragraphs {idx_mq_recovery}..{idx_mq_recovery+2})")

# ---------- 8. Update 4.2.3.2 订单簿数据结构 ----------
idx_orderbook_intro = find_para_index('订单簿采用"两层结构"')
assert idx_orderbook_intro >= 0, "Could not find 订单簿采用 两层结构"

new_intro = (
    '订单簿统一采用 sortedcontainers.SortedDict + PriceLevel + deque + order_index：'
    'SortedDict 只负责价格档位排序；同一价格档位内使用 deque 按入队时间 FIFO 排列；order_index 用于撤单时按 order_id 快速定位。'
    '订单簿整体为"两层结构"：第一层按价格档位（PriceLevel）组织，第二层在同一价位内按时间顺序组织订单。'
)
doc.paragraphs[idx_orderbook_intro].text = new_intro
print(f"[OK] Updated 4.2.3.2 订单簿数据结构 (paragraph {idx_orderbook_intro})")

# Remove the blank separator paragraph that followed
blank_p = doc.paragraphs[idx_orderbook_intro + 1]
if not blank_p.text.strip():
    blank_p._element.getparent().remove(blank_p._element)
    print(f"[OK] Removed blank separator after order book intro")
else:
    print(f"[INFO] Paragraph after intro not blank, skipping: {blank_p.text[:50]!r}")

# ---------- 9. Collapse redundant empty paragraphs in the cover area ----------
# We want at most 1 empty paragraph in a row in the cover area (before first heading).
# Find the first heading and stop there.
removed = 0
prev_empty = None
runs = 0
in_cover = True
for p in list(doc.paragraphs):
    if not in_cover:
        break
    if p.style.name.startswith('Heading'):
        in_cover = False
        break
    if not p.text.strip():
        runs += 1
        if runs > 1 and prev_empty is not None:
            try:
                prev_empty._element.getparent().remove(prev_empty._element)
                removed += 1
                # Don't update prev_empty to the removed one
            except Exception:
                pass
        else:
            prev_empty = p
    else:
        runs = 0
        prev_empty = None

print(f"[OK] Collapsed {removed} redundant empty paragraphs on cover page")

doc.save(DST)
print(f"\n[SAVED] {DST}")
